import os
from pathlib import Path
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from src.watermark import create_transparent_watermark


def add_page_number(footer):
    """
    Adds a dynamic MS Word page number field centered in the footer.
    """
    # Check if page number already exists in footer
    footer_xml = footer._element.xml
    if "w:fldSimple" in footer_xml or "PAGE" in footer_xml:
        return

    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add page field XML element
    fldSimple = parse_xml(r'<w:fldSimple %s w:instr="PAGE"/>' % nsdecls('w'))
    p._p.append(fldSimple)


def add_top_right_logo_vml(header, logo_path: str, width_inches: float = 2.0, height_inches: float = 2.0, top_offset_in: float = 0.5, right_offset_in: float = 0.5):
    """
    Inserts 2" x 2" top-right logo into header, positioned 0.5" from top/right corner behind text.
    """
    header_xml = header._element.xml
    if "TopRightLogoShape" in header_xml:
        return

    p = header.add_paragraph() if header.paragraphs else header.add_paragraph()
    run = p.add_run()
    
    inline = run.add_picture(logo_path, width=Inches(width_inches), height=Inches(height_inches))
    r_id = inline._inline.graphic.graphicData.pic.blipFill.blip.embed

    left_pt = int((8.5 - right_offset_in - width_inches) * 72)
    top_pt = int(top_offset_in * 72)
    width_pt = int(width_inches * 72)
    height_pt = int(height_inches * 72)

    vml_xml = f'''
    <w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            xmlns:v="urn:schemas-microsoft-com:vml"
            xmlns:o="urn:schemas-microsoft-com:office:office"
            xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
      <v:shape id="TopRightLogoShape"
               style="position:absolute;left:{left_pt}pt;top:{top_pt}pt;width:{width_pt}pt;height:{height_pt}pt;z-index:-251658239;mso-position-horizontal:right;mso-position-horizontal-relative:page;mso-position-vertical:top;mso-position-vertical-relative:page;margin-right:36pt"
               type="#_x0000_t75">
        <v:imagedata r:id="{r_id}" o:title="HeaderLogo"/>
      </v:shape>
    </w:pict>
    '''
    pict = parse_xml(vml_xml)
    p._p.remove(run._r)
    p_logo = header.add_paragraph()
    p_logo._p.append(pict)


def add_center_watermark_vml(header, transparent_logo_path: str):
    """
    Inserts a centered background watermark image into section header using VML behind text.
    The logo is pre-processed with 93% transparency (7% opacity).
    """
    header_xml = header._element.xml
    if "CenterWatermark" in header_xml:
        return

    p = header.add_paragraph()
    run = p.add_run()
    
    inline = run.add_picture(transparent_logo_path, width=Inches(5.5), height=Inches(5.5))
    r_id = inline._inline.graphic.graphicData.pic.blipFill.blip.embed

    vml_xml = f'''
    <w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            xmlns:v="urn:schemas-microsoft-com:vml"
            xmlns:o="urn:schemas-microsoft-com:office:office"
            xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
      <v:shape id="CenterWatermark"
               style="position:absolute;margin-left:0;margin-top:0;width:450pt;height:450pt;z-index:-251658240;mso-position-horizontal:center;mso-position-horizontal-relative:page;mso-position-vertical:center;mso-position-vertical-relative:page"
               type="#_x0000_t75">
        <v:imagedata r:id="{r_id}" o:title="Watermark"/>
      </v:shape>
    </w:pict>
    '''
    pict = parse_xml(vml_xml)

    p._p.remove(run._r)
    p_wm = header.add_paragraph()
    p_wm._p.append(pict)


def format_existing_docx(
    input_docx_path: str,
    cover_image_path: str,
    logo_path: str,
    output_docx_path: str,
    top_right_logo_width: float = 2.0,
    top_right_logo_height: float = 2.0,
    transparency_percent: float = 93.0
) -> str:
    """
    Formats an existing .docx file directly:
    - 100% PRESERVES all existing photos, images, tables, drawings, and formatting inside the document.
    - Prepends the Cover Page (with cover photo) on Page 1.
    - Applies Letter size (8.5" x 11.0").
    - Adds 2"x2" top-right corner logo (behind text, 0.5" offset).
    - Adds 93% transparent center watermark (behind text).
    - Adds bottom-center page numbers.
    """
    doc = Document(str(input_docx_path))

    # Pre-process logo to 93% transparent (7% opacity)
    transparent_logo_path = create_transparent_watermark(
        image_path=logo_path,
        transparency_percent=transparency_percent
    )

    # Insert cover page at the beginning if cover image is provided
    if cover_image_path and Path(cover_image_path).exists():
        if doc.paragraphs:
            first_p = doc.paragraphs[0]
            cover_p = first_p.insert_paragraph_before()
            cover_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cover_run = cover_p.add_run()
            cover_run.add_picture(str(cover_image_path), width=Inches(6.5))

            # Page break after cover
            break_p = first_p.insert_paragraph_before()
            break_p.add_run().add_break(docx.enum.text.WD_BREAK.PAGE)
        else:
            cover_p = doc.add_paragraph()
            cover_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cover_run = cover_p.add_run()
            cover_run.add_picture(str(cover_image_path), width=Inches(6.5))
            doc.add_page_break()

    # Format sections
    for i, section in enumerate(doc.sections):
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

        # First section has cover page on page 1
        if i == 0:
            section.different_first_page_header_footer = True

        # Apply header (logo + watermark)
        header = section.header
        add_top_right_logo_vml(
            header=header,
            logo_path=logo_path,
            width_inches=top_right_logo_width,
            height_inches=top_right_logo_height,
            top_offset_in=0.5,
            right_offset_in=0.5
        )
        add_center_watermark_vml(
            header=header,
            transparent_logo_path=transparent_logo_path
        )

        # Apply footer (page numbers)
        footer = section.footer
        add_page_number(footer=footer)

    output_path = Path(output_docx_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return str(output_path)


def build_docx_document(
    cover_image_path: str,
    content_elements: list[dict],
    logo_path: str,
    output_docx_path: str,
    top_right_logo_width: float = 2.0,
    top_right_logo_height: float = 2.0,
    transparency_percent: float = 93.0
) -> str:
    """
    Builds Microsoft Word (.docx) document from structured content elements:
    1. Letter size 8.5" x 11.0"
    2. Cover page (extracted screenshot cover image)
    3. Header with 2" x 2" top-right logo positioned 0.5" from top/right corner behind text
    4. Center 93% transparent watermark logo behind text
    5. Footer with page number centered at bottom line
    """
    doc = Document()

    transparent_logo_path = create_transparent_watermark(
        image_path=logo_path,
        transparency_percent=transparency_percent
    )

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11.0)
    
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    
    section.different_first_page_header_footer = True

    # --- 1. COVER PAGE (First Page) ---
    if cover_image_path and Path(cover_image_path).exists():
        cover_p = doc.add_paragraph()
        cover_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cover_run = cover_p.add_run()
        cover_run.add_picture(str(cover_image_path), width=Inches(6.5))
        doc.add_page_break()

    # --- 2. BODY HEADERS & FOOTERS ---
    header = section.header
    footer = section.footer

    add_top_right_logo_vml(
        header=header,
        logo_path=logo_path,
        width_inches=top_right_logo_width,
        height_inches=top_right_logo_height,
        top_offset_in=0.5,
        right_offset_in=0.5
    )

    add_center_watermark_vml(
        header=header,
        transparent_logo_path=transparent_logo_path
    )

    add_page_number(footer=footer)

    # --- 3. ADD BODY CONTENT ---
    for elem in content_elements:
        elem_type = elem.get("type")
        text = elem.get("text", "")
        
        if elem_type == "heading":
            level = elem.get("level", 1)
            h = doc.add_heading(text, level=min(level, 3))
            h.paragraph_format.space_before = Pt(14)
            h.paragraph_format.space_after = Pt(6)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(text)
            run.font.name = "Arial"
            run.font.size = Pt(11)

    output_path = Path(output_docx_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    return str(output_path)
