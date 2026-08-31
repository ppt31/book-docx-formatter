import os
from pathlib import Path
from PIL import Image, ImageDraw
import docx
from docx import Document
from docx.shared import Inches

BASE_DIR = Path(__file__).resolve().parent.parent
INPUT_DIR = BASE_DIR / "input_books"
OUTPUT_DIR = BASE_DIR / "output_docx"

def create_docx_with_photos():
    docx_path = INPUT_DIR / "sample_book_with_photos.docx"
    doc = Document()
    
    doc.add_heading("Chapter 1: Book with Embedded Photos", level=1)
    doc.add_paragraph("This is a test paragraph before the photo.")
    
    # Create and add a photo
    photo_path = INPUT_DIR / "test_photo.png"
    img = Image.new("RGB", (400, 300), color=(200, 100, 50))
    d = ImageDraw.Draw(img)
    d.text((50, 130), "EMBEDDED BOOK PHOTO", fill=(255, 255, 255))
    img.save(photo_path)
    
    # Add photo to docx
    doc.add_picture(str(photo_path), width=Inches(4.0))
    doc.add_paragraph("This is text after the embedded photo.")
    
    doc.save(str(docx_path))
    print(f"[+] Created DOCX with embedded photo at: {docx_path}")
    return docx_path

if __name__ == "__main__":
    create_docx_with_photos()
