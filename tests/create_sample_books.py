import os
import io
from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
import fitz  # PyMuPDF
import ebooklib
from ebooklib import epub
import docx

BASE_DIR = Path(__file__).resolve().parent.parent
INPUT_DIR = BASE_DIR / "input_books"
INPUT_DIR.mkdir(parents=True, exist_ok=True)


def create_sample_pdf():
    pdf_path = INPUT_DIR / "sample_english_burmese.pdf"
    doc = fitz.open()

    # Page 0: Cover Page
    page_cover = doc.new_page(width=595, height=842)  # A4 size
    rect = fitz.Rect(0, 0, 595, 842)
    shape = page_cover.new_shape()
    shape.draw_rect(rect)
    shape.finish(fill=(0.1, 0.2, 0.4), color=None)
    shape.commit()

    page_cover.insert_text(
        fitz.Point(60, 250), 
        "ENGLISH TO BURMESE\nTRANSLATION BOOK", 
        fontsize=28, 
        color=(1, 1, 1)
    )
    page_cover.insert_text(
        fitz.Point(60, 350), 
        "Bilingual Edition / အင်္ဂလိပ် - မြန်မာ ဘာသာပြန် မာတိကာ", 
        fontsize=16, 
        color=(0.9, 0.9, 0.9)
    )

    # Page 1: Content
    page_body = doc.new_page(width=595, height=842)
    body_text = """
Chapter 1: Daily Conversations (အခန်း ၁ - နေ့စဉ်ပြောစကားများ)

1. Hello! How are you today?
   မင်္ဂလာပါ! ဒီနေ့ နေကောင်းရဲ့လား။

2. I am learning English and Burmese languages.
   ကျွန်တော် အင်္ဂလိပ်စာနှင့် မြန်မာစာကို လေ့လာနေပါတယ်။

3. Thank you very much for reading this translation book.
   ဤဘာသာပြန်စာအုပ်ကို ဖတ်ရှုပေးသည့်အတွက် အထူးကျေးဇူးတင်ရှိပါသည်။
"""
    page_body.insert_text(fitz.Point(50, 80), body_text, fontsize=13, color=(0, 0, 0))

    doc.save(str(pdf_path))
    doc.close()
    print(f"[+] Created sample PDF book at: {pdf_path}")
    return pdf_path


def create_sample_epub():
    epub_path = INPUT_DIR / "sample_english_burmese.epub"
    book = epub.EpubBook()

    book.set_identifier("eng-bur-sample-001")
    book.set_title("English to Burmese Translation Handbook")
    book.set_language("en")
    book.add_author("Antigravity Translation Press")

    cover_img = Image.new("RGB", (600, 800), color=(30, 80, 160))
    draw = ImageDraw.Draw(cover_img)
    draw.text((80, 300), "ENGLISH - BURMESE\nHANDBOOK", fill=(255, 255, 255), font_size=40)
    
    img_byte_arr = io.BytesIO()
    cover_img.save(img_byte_arr, format="PNG")
    cover_bytes = img_byte_arr.getvalue()

    book.set_cover("cover.png", cover_bytes)

    ch1 = epub.EpubHtml(title="Chapter 1", file_name="ch1.xhtml", lang="en")
    ch1.content = """
    <h1>Chapter 1: Greetings & Polite Expressions</h1>
    <p><b>English:</b> Good morning! Welcome to the translation guide.</p>
    <p><b>Burmese:</b> မင်္ဂလာနံနက်ခင်းပါ! ဘာသာပြန် လမ်းညွှန်မှ ကြိုဆိုပါတယ်။</p>
    """
    book.add_item(ch1)
    book.spine = ["nav", ch1]

    epub.write_epub(str(epub_path), book)
    print(f"[+] Created sample EPUB book at: {epub_path}")
    return epub_path


def create_sample_docx():
    docx_path = INPUT_DIR / "sample_english_burmese.docx"
    doc = docx.Document()

    doc.add_heading("English to Burmese Translation Book", level=1)
    doc.add_paragraph("Chapter 1: Essential Phrases (အခန်း ၁ - အဓိက စကားစုများ)")
    doc.add_paragraph("English: Could you please help me translate this document?")
    doc.add_paragraph("Burmese: ဤစာရွက်စာတမ်းကို ဘာသာပြန်ဆိုရန် ကူညီပေးနိုင်မလား။")

    doc.save(str(docx_path))
    print(f"[+] Created sample DOCX book at: {docx_path}")
    return docx_path


if __name__ == "__main__":
    create_sample_pdf()
    create_sample_epub()
    create_sample_docx()
